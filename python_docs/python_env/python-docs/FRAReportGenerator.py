import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocumentGenerator:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Document Generator")

        self.create_document_type_page()

    def create_document_type_page(self):
        self.document_type_var = tk.StringVar()

        label = tk.Label(self.root, text="Document type:")
        label.pack(pady=10)

        document_type_options = ["Notes", "Causation report", "Rebuttal", "Affidavit"]
        document_type_dropdown = ttk.Combobox(self.root, values=document_type_options, textvariable=self.document_type_var)
        document_type_dropdown.pack(pady=10)

        next_button = tk.Button(self.root, text="Next", command=self.show_attorney_info_page)
        next_button.pack(pady=10)

    def show_attorney_info_page(self):
        document_type = self.document_type_var.get()

        if not document_type:
            tk.messagebox.showerror("Error", "Please select a document type.")
            return

        # Create a new top-level window for attorney info
        attorney_info_window = tk.Toplevel(self.root)
        attorney_info_window.title("Attorney Information")

        # Create attorney info fields
        self.attorney_first_name_var = tk.StringVar()
        self.attorney_last_name_var = tk.StringVar()
        self.attorney_gender_var = tk.StringVar()

        label_first_name = tk.Label(attorney_info_window, text="Attorney first name:")
        label_first_name.pack(pady=10)
        entry_first_name = tk.Entry(attorney_info_window, textvariable=self.attorney_first_name_var)
        entry_first_name.pack(pady=10)

        label_last_name = tk.Label(attorney_info_window, text="Attorney last name:")
        label_last_name.pack(pady=10)
        entry_last_name = tk.Entry(attorney_info_window, textvariable=self.attorney_last_name_var)
        entry_last_name.pack(pady=10)

        label_gender = tk.Label(attorney_info_window, text="Attorney gender:")
        label_gender.pack(pady=10)
        gender_options = ["Man", "Woman"]
        gender_dropdown = ttk.Combobox(attorney_info_window, values=gender_options, textvariable=self.attorney_gender_var)
        gender_dropdown.pack(pady=10)

        # Add an additional field for "Case number" if the document type is "Causation report"
        if document_type == "Causation report":
            self.case_number_var = tk.StringVar()
            label_case_number = tk.Label(attorney_info_window, text="Case number:")
            label_case_number.pack(pady=10)
            entry_case_number = tk.Entry(attorney_info_window, textvariable=self.case_number_var)
            entry_case_number.pack(pady=10)

        compile_button = tk.Button(attorney_info_window, text="Compile", command=lambda: self.compile_document(attorney_info_window))
        compile_button.pack(pady=10)

    def compile_document(self, attorney_info_window):
        first_name = self.attorney_first_name_var.get()
        last_name = self.attorney_last_name_var.get()
        gender = self.attorney_gender_var.get()

        if not all([first_name, last_name, gender]):
            tk.messagebox.showerror("Error", "Please fill in all the empty boxes.")
            return

        # Create a new Word document
        doc = Document()

        # Add attorney name to the document
        attorney_name = f"{first_name} {last_name}"
        doc.add_paragraph(attorney_name)

        # Add a line break
        doc.add_paragraph()

        # Add gender information to the document
        doc.add_paragraph(f"Gender: {gender.lower()}")

        # Add case number if available
        if hasattr(self, 'case_number_var'):
            case_number = self.case_number_var.get()
            doc.add_paragraph(f"Case number: {case_number}")

        # Save the document
        file_path = f"{first_name}_{last_name}_document.docx"
        doc.save(file_path)

        # Close the attorney info window
        attorney_info_window.destroy()

        print(f"Document '{file_path}' created successfully.")

if __name__ == "__main__":
    doc_generator = DocumentGenerator()
    doc_generator.root.mainloop()